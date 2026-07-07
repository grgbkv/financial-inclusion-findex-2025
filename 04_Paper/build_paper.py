"""Build the working paper .docx from paper_stats.json and figures_print/.

Run from 04_Paper/:  python3 build_paper.py
"""
import json
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
S = json.load(open(os.path.join(HERE, 'paper_stats.json')))
FIG = os.path.join(HERE, 'figures_print')

doc = Document()

# ---------------------------------------------------------------- page & base style
for sec in doc.sections:
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, m, Inches(1))

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.25

# page number in footer
footer_p = doc.sections[0].footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run()
for el, attrs, text in [('w:fldChar', {'w:fldCharType': 'begin'}, None),
                        ('w:instrText', {'xml:space': 'preserve'}, ' PAGE '),
                        ('w:fldChar', {'w:fldCharType': 'end'}, None)]:
    e = OxmlElement(el)
    for k, v in attrs.items():
        e.set(qn(k) if ':' in k else k, v)
    if text:
        e.text = text
    run._r.append(e)
footer_p.runs[0].font.size = Pt(10)


def para(text, size=12, bold=False, italic=False, align='justify', space_after=6,
         space_before=0, color=None, keep_next=False):
    p = doc.add_paragraph()
    p.alignment = {'justify': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'left': WD_ALIGN_PARAGRAPH.LEFT}[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if keep_next:
        p.paragraph_format.keep_with_next = True
    # minimal markdown: *italic* segments
    parts = text.split('*')
    for i, chunk in enumerate(parts):
        if not chunk:
            continue
        r = p.add_run(chunk)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic or (i % 2 == 1)
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p


def heading(num, text, level=1):
    label = f'{num}. {text}' if num else text
    p = para(label, size=13 if level == 1 else 12, bold=True, align='left',
             space_after=6, space_before=14 if level == 1 else 10, keep_next=True)
    return p


def caption(text):
    para(text, size=9.5, align='justify', space_after=10)


def figure(path, caption_text=None, width=5.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(os.path.join(FIG, path), width=Inches(width))
    caption(caption_text)


def set_cell(cell, text, bold=False, size=10.5, align='left'):
    cell.paragraphs[0].text = ''
    p = cell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                   'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)


def make_table(header, rows, col_widths=None, caption_above=None, cell_size=10.5):
    if caption_above:
        para(caption_above, size=10, bold=False, align='left', space_after=4, space_before=8,
             keep_next=True)
    t = doc.add_table(rows=len(rows) + 1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for j, h in enumerate(header):
        set_cell(t.rows[0].cells[j], h, bold=True, size=cell_size, align='center' if j else 'left')
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            set_cell(t.rows[i].cells[j], str(v), size=cell_size, align='center' if j else 'left')
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    # keep the whole table on one page
    for row in t.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
    for row in list(t.rows)[:-1]:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_with_next = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


Y = [2011, 2014, 2017, 2021, 2024]
g = lambda d, y: d[str(y)] if str(y) in d else d[y]  # json keys are strings

# ================================================================ TITLE PAGE
para('', space_after=30)
para('Access Without Depth: Financial Inclusion 2011–2024 and Aggregation '
     'Pitfalls in the Global Findex Database 2025',
     size=17, bold=True, align='center', space_after=18)
para('[Author Name]', size=13, align='center', space_after=2)
para('[Affiliation]', size=12, align='center', space_after=2)
para('[email@example.com]', size=11, align='center', space_after=24)
para('Working Paper — July 2026', size=11, italic=True, align='center', space_after=28)

para('Abstract', size=12, bold=True, align='center', space_after=8)
para('The Global Findex Database 2025 extends the World Bank’s survey of how adults access and '
     'use financial services to a fifth wave. Using a balanced panel of 117 economies covering '
     '96.5 percent of the surveyed adult population, this paper documents the corrected long-run '
     'picture: account ownership rose from 51 to 79 percent of adults between 2011 and 2024 — '
     'nearly two billion adults gained access — while formal saving in developing economies, flat '
     'for a decade, surged by 14 percentage points between 2021 and 2024, largely over mobile-money '
     'rails. Financial resilience, by contrast, stagnated near 55 percent, and the income gap in '
     'account ownership (10.5 points) remains wider than the gender gap, which halved to 4 points. '
     'The paper’s second contribution is methodological. The Findex 2025 country file stacks '
     'demographic disaggregations, embeds the World Bank’s own aggregate rows, and varies '
     'indicator coverage across waves. Four resulting aggregation pitfalls are quantified. '
     'Individually each biases levels by only one to two percentage points; compounded, they '
     'reverse trend signs — turning stable resilience into an apparent decline, rising borrowing '
     'into an apparent fall, and almost doubling the global mobile-money figure. A validation-first '
     'workflow that benchmarks every computed series against the embedded official aggregates '
     'catches all four. Code and data are openly available.',
     size=11, align='justify', space_after=12)
para('*Keywords:* financial inclusion; Global Findex; mobile money; digital payments; '
     'measurement error; reproducibility', size=11, align='left', space_after=4)
para('*JEL codes:* C81, G21, G51, O16', size=11, align='left', space_after=4)
para('*Data and code:* [repository URL]', size=11, align='left', space_after=4)
para('This paper grew out of the author’s bachelor’s thesis; the empirical analysis was '
     'rebuilt and audited for this version, and the audit of the earlier version supplied the '
     'error catalogue analysed in Section 5. All remaining errors are the author’s own.',
     size=10, italic=True, align='justify', space_before=16)
doc.add_page_break()

# ================================================================ 1 INTRODUCTION
heading(1, 'Introduction')
para('Access to formal financial services — an account to receive wages, a safe place to save, a '
     'channel to pay and be paid — is a precondition for participating in a modern economy and an '
     'enabler of several Sustainable Development Goals. Since 2011 the World Bank’s Global '
     'Findex Database has been the standard yardstick for measuring that access: a nationally '
     'representative survey of adults in more than 140 economies, repeated roughly every three '
     'years (Demirgüç-Kunt and Klapper, 2012). The 2025 edition adds a fifth wave, '
     'collected in 2024, and with it the first full picture of how financial inclusion evolved '
     'through the pandemic period and its aftermath (Klapper, Singer, Starita and Norris, 2025).')
para('This paper does two things. First, it documents the headline trends of 2011–2024 on a '
     'balanced panel of 117 economies — the same set of countries in every wave — so that changes '
     'over time reflect behaviour rather than a shifting country mix. The corrected picture is '
     'more interesting than a simple success story: access expanded enormously (51 to 79 percent '
     'of adults) and formal saving finally surged in the latest wave (+14 percentage points in '
     'developing economies, the largest wave-on-wave change in the series), but financial '
     'resilience — the ability to raise emergency funds — barely moved, and the income gap in '
     'account ownership remains wider than the gender gap. Access, in short, has not yet '
     'translated into depth.')
para('Second, the paper treats the dataset itself as an object of study. The Findex 2025 '
     'country-level file has three structural features that make naive aggregation hazardous: it '
     'stacks demographic disaggregations (gender, age, income, education, labour-force status) '
     'as additional rows beneath each economy’s total-population row; it embeds the World '
     'Bank’s own regional, income-group and world aggregates as pseudo-economy rows in the '
     'same file; and its indicator coverage varies across waves and economies — mobile money is '
     'only surveyed where such services operate, and the digital-payment composite was collected '
     'for just five of forty high-income panel economies in 2024. Section 5 formalises four '
     'resulting pitfalls and quantifies each. The striking result is not that mistakes are '
     'possible — it is their arithmetic: each pitfall alone shifts levels by only one to two '
     'percentage points, easily dismissed as noise, but plausible combinations reverse the sign '
     'of headline trends. A stable resilience series becomes a four-point decline; rising formal '
     'borrowing becomes a fall; the global mobile-money share nearly doubles.')
para('The economics literature has a well-known precedent for consequential aggregation error: '
     'the spreadsheet and weighting problems that overturned a widely cited result on public debt '
     'and growth (Herndon, Ash and Pollin, 2014). The pitfalls documented here are smaller in '
     'consequence but far easier to fall into, because the Findex file’s structure invites '
     'them. The remedy proposed is procedural rather than clever: because the publisher ships its '
     'own aggregates inside the file, every computed series can be benchmarked against an official '
     'counterpart before any substantive conclusion is drawn. This validation-first workflow, '
     'applied throughout the paper (every figure carries the official aggregate as hollow '
     'markers), catches all four pitfalls immediately. The full pipeline — from raw file to every '
     'figure and table — is open and reproducible.')
para('The rest of the paper proceeds as follows. Section 2 describes the data. Section 3 sets out '
     'the sample, weighting and validation methodology. Section 4 presents the corrected trends '
     'in access, usage, resilience and equality of access. Section 5 quantifies the aggregation '
     'pitfalls. Section 6 discusses limitations, and Section 7 concludes.')

# ================================================================ 2 DATA
heading(2, 'Data')
para('The Global Findex Database 2025 is built from nationally representative surveys of the '
     'population aged 15 and older, carried out as part of the Gallup World Poll with roughly '
     '1,000 respondents per economy per wave (World Bank, 2025). The country-level file used here '
     'covers 162 economies over five waves — 2011, 2014, 2017, 2021 and 2024 — with 139 to 146 '
     'economies per wave. Sixteen economies whose 2021-round fieldwork was delayed by COVID-19 '
     'restrictions appear in the file under a nominal “2022” wave; following World Bank '
     'practice, they are merged into the 2021 wave here, before any other processing.')
para('The file is in long format: one row per economy, wave and population group, with 437 '
     'indicator columns expressed as population shares. The group dimension stacks the total-adult '
     'row (group “all”) together with ten to twelve demographic subgroups per economy '
     'and wave. Interspersed with the economies are twelve aggregate entities published by the '
     'World Bank itself — the world, developing economies as a block, seven regional groupings '
     'and four income groups. These aggregate rows carry no adult-population weight and no region '
     'code, which is what distinguishes them mechanically from economy rows. Both structural '
     'facts — stacked subgroups and embedded aggregates — are central to Section 5.')
para('Indicators are referenced by their database codes. Table 1 maps the concepts used in this '
     'paper to those codes. Two definitional details matter. The headline saving and borrowing '
     'measures (“saved formally”, “borrowed formally”) include activity '
     'through mobile-money accounts, not only at banks and similar institutions; their narrower '
     '“at a financial institution” variants coexist in the file under similar names, '
     'and the two families diverge sharply after 2021. Account dormancy is expressed as a share '
     'of adults and is related to account ownership below to obtain the share of account holders '
     'whose account is inactive.')

make_table(
    ['Concept', 'Definition (share of adults 15+)', 'Findex code'],
    [
        ['Account ownership', 'Has an account at a financial institution or mobile-money provider', 'account_t_d'],
        ['Mobile money account', 'Has a registered mobile-money account', 'mobileaccount_t_d'],
        ['Digital payment', 'Made or received a digital payment in the past year', 'g20_any'],
        ['Saved formally', 'Saved at a financial institution or via a mobile-money account', 'fin17a_17a1_d'],
        ['Borrowed formally', 'Borrowed from a financial institution, credit card, or mobile-money account', 'fin22a_22a1_22g_d'],
        ['Inactive account', 'Has an account with no deposit or withdrawal in the past year', 'inactive_t_d'],
        ['Emergency funds', 'Could raise emergency funds with some or no difficulty', 'fin24aSD_ND'],
    ],
    col_widths=[1.55, 3.75, 1.2],
    caption_above='*Table 1.* Concepts and Findex 2025 indicator codes used in this paper.')

# ================================================================ 3 METHODOLOGY
heading(3, 'Methodology')
heading('3.1', 'Balanced panel', level=2)
para(f'All headline series are computed on the {S["n_panel"]} economies present in every wave '
     f'({S["n_panel_dev"]} of them developing). The panel restriction costs little coverage — '
     f'{S["panel_coverage_2024"]} percent of the adult population of all surveyed economies in '
     '2024 — and buys compositional consistency: without it, part of any measured change is '
     'simply the set of surveyed economies changing between waves. Section 5 shows this is not a '
     'theoretical concern.')
heading('3.2', 'Population-weighted aggregation', level=2)
para('Country indicators are aggregated with adult-population weights, matching World Bank '
     'methodology: an aggregate share is Σ wᵢxᵢ / Σ wᵢ, where xᵢ is '
     'economy i’s indicator value and wᵢ its adult population in the given wave. Only '
     'total-population rows (group “all”) of economy entities enter these calculations; '
     'gender- and income-group series use the corresponding subgroup rows, one row per economy '
     'and wave. Ratio measures (the dormancy ratio; ownership gaps) are formed from the weighted '
     'aggregates, not averaged across countries.')
heading('3.3', 'Validation against embedded official aggregates', level=2)
para('Because the file ships the publisher’s own aggregates, every computed series can be '
     'benchmarked before interpretation. Table 2 reports the comparison for the headline '
     'indicator. The panel tracks the official world series to within '
     f'{S["account_maxdev"]} percentage points, and the deviation has a mechanical explanation: '
     'the World Bank imputes values for economies missing from a wave, and those economies are '
     'disproportionately poorer, so the official series sits slightly below the panel. Throughout '
     'the paper, figures carry the official aggregates as hollow markers alongside the panel '
     'estimates.')

make_table(
    ['Wave', 'Balanced panel', 'Official world aggregate', 'Difference (pp)'],
    [[y, f'{g(S["account_panel"], y):.1f}', f'{g(S["account_official"], y):.1f}',
      f'{g(S["account_panel"], y) - g(S["account_official"], y):+.1f}'] for y in Y],
    col_widths=[0.9, 1.9, 2.3, 1.4],
    caption_above='*Table 2.* Account ownership, percent of adults: balanced panel vs. the official '
                  'world aggregate embedded in the Findex 2025 file.')

heading('3.4', 'A taxonomy of aggregation pitfalls', level=2)
para('Section 5 quantifies four pitfalls, each mapped to a structural feature of the file. '
     '*P1 — disaggregation-row leakage:* aggregating without fixing the group dimension, so that '
     'every demographic subgroup row enters the average alongside the total-population row, each '
     'carrying the full country weight. *P2 — unbalanced composition:* aggregating over whichever '
     'economies happen to be surveyed in each wave. *P3 — indicator-variant confusion:* using a '
     'narrower or broader variant of an indicator family than the concept requires. '
     '*P4 — coverage-driven missingness:* averaging only the economies that report an indicator '
     'when non-reporting is informative (mobile money is not surveyed where services are absent) '
     'or when coverage collapses in a wave (the high-income digital-payment composite in 2024).')

# ================================================================ 4 RESULTS
heading(4, 'Results')
heading('4.1', 'Access: the decade of the account', level=2)
para(f'Account ownership rose from {g(S["account_panel"],2011):.1f} percent of adults in 2011 to '
     f'{g(S["account_panel"],2024):.1f} percent in 2024 (Figure 1) — in absolute terms, from '
     f'{S["adults_with_accounts_bn"]["2011"]:.1f} to {S["adults_with_accounts_bn"]["2024"]:.1f} '
     'billion adults on the panel, an expansion of nearly two billion people. Growth decelerates '
     'as saturation approaches: +11.4 points in 2011–2014, +3.6 points in 2021–2024.', keep_next=True)
figure('fig1_account.png', width=4.7, caption_text=
       '*Figure 1.* Account ownership, share of adults, 2011–2024. Solid line: balanced panel of '
       '117 economies; hollow markers: official world aggregate. Source: author’s '
       'calculations from the Global Findex Database 2025.')
para('The regional picture (Figure 2) contains three different convergence stories. South Asia '
     f'made the largest leap ({g(S["regional_account"]["South Asia"],2011):.1f} to '
     f'{g(S["regional_account"]["South Asia"],2024):.1f} percent), concentrated in the 2014–2017 '
     'period of India’s mass account-opening programme. Sub-Saharan Africa more than doubled '
     f'from the lowest base ({g(S["regional_account"]["Sub-Saharan Africa"],2011):.1f} to '
     f'{g(S["regional_account"]["Sub-Saharan Africa"],2024):.1f} percent). Latin America climbed '
     f'steadily to {g(S["regional_account"]["Latin America & Caribbean"],2024):.1f} percent. '
     'High-income economies, near saturation throughout, drifted from '
     f'{g(S["regional_account"]["High income"],2011):.1f} to '
     f'{g(S["regional_account"]["High income"],2024):.1f} percent.', keep_next=True)
figure('fig2_regional.png', width=5.0,
       caption_text='*Figure 2.* Account ownership by region. Hollow markers: official regional aggregates. '
       'Source: author’s calculations from the Global Findex Database 2025.')

heading('4.2', 'Mobile money: Africa’s parallel banking system', level=2)
para('Mobile money is the indicator where measurement care matters most, because it is only '
     'surveyed where such services operate. Treating non-surveyed economies as effectively '
     'zero-adoption reproduces the official world aggregate almost exactly (Section 5, pitfall '
     'P4).')
para(f'On that basis, global ownership reached {g(S["mm_filled"],2024):.1f} percent of '
     f'adults in 2024 (official: {g(S["mm_official"],2024):.1f}). The global figure is, however, '
     'mostly a Sub-Saharan African story (Figure 3): there, mobile-money ownership reached '
     f'{g(S["mm_ssa"],2024):.1f} percent of adults — exceeding ownership of accounts at financial '
     f'institutions ({S["ssa_fiaccount_2024"]:.1f} percent). For much of the region the phone, '
     'not the bank branch, is the financial system’s front door (Suri and Jack, 2016).', keep_next=True)
figure('fig3_mobile_money.png', width=4.7, caption_text=
       '*Figure 3.* Mobile-money account ownership, 2014–2024. World series treats non-surveyed '
       'economies as zero-adoption; hollow markers: official world aggregate. Mobile money was '
       'not measured in 2011. Source: author’s calculations from the Global Findex Database '
       '2025.')

heading('4.3', 'Usage: payments first, saving finally, borrowing slowly', level=2)
para('Digital payments are the fastest-moving usage indicator: the share of adults in developing '
     f'economies who made or received one rose from {g(S["dp_dev"],2014):.1f} percent in 2014 to '
     f'{g(S["dp_dev"],2024):.1f} percent in 2024. Regionally (Figure 4), Sub-Saharan Africa’s '
     'curve rides on mobile money while South Asia’s reflects public digital-payment '
     'infrastructure. The high-income series ends in 2021: the composite was collected for only '
     f'{g(S["hi_g20_coverage"],2024)} of {S["n_hi_panel"]} high-income panel economies in the '
     '2024 wave, too few to represent the group (Section 5, pitfall P4).', keep_next=True)
figure('fig4_dp_regional.png', width=5.0,
       caption_text='*Figure 4.* Made or received a digital payment, by region. The high-income line ends in '
       '2021 because the indicator was collected for only 5 of 40 high-income panel economies in '
       '2024. Source: author’s calculations from the Global Findex Database 2025.')
para('The flagship change of the 2025 wave is in saving (Figure 5). Formal saving in developing '
     'economies was essentially flat for a decade — '
     f'{g(S["sav_dev"],2011):.1f} percent in 2011, {g(S["sav_dev"],2021):.1f} percent in 2021 — '
     f'and then jumped to {g(S["sav_dev"],2024):.1f} percent in 2024, a rise of '
     f'{g(S["sav_dev"],2024)-g(S["sav_dev"],2021):.1f} points on the panel (the official '
     f'aggregate puts it at +16.0 points, to {g(S["sav_dev_official"],2024):.1f} percent).')
para('Decomposing the headline definition shows mobile-money accounts contribute '
     f'{g(S["sav_dev"],2024)-g(S["sav_narrow_dev"],2024):.1f} points of the 2024 level — the '
     'saving surge largely travelled on the rails built by the payments expansion. Formal '
     f'borrowing, by contrast, moved gradually throughout: {g(S["bor_dev"],2014):.1f} percent in '
     f'2014 to {g(S["bor_dev"],2024):.1f} percent in 2024, rising, not falling (a point that '
     'matters in Section 5).', keep_next=True)
figure('fig5_sav_bor.png', width=5.0,
       caption_text='*Figure 5.* Saved formally and borrowed formally, developing economies. Hollow markers: '
       'official aggregates for developing economies. Source: author’s calculations from the '
       'Global Findex Database 2025.')
para('Dormancy — the quality check on the access story — peaked at '
     f'{g(S["inactivity_ratio"],2017):.1f} percent of account holders in 2017, immediately after '
     'the largest account-opening drives, and fell to '
     f'{g(S["inactivity_ratio"],2024):.1f} percent by 2024. Newly opened accounts increasingly '
     'get used.')

heading('4.4', 'Resilience: the dimension that did not move', level=2)
para('The ability to raise emergency funds with at most some difficulty — the survey’s '
     'resilience proxy, available for 2021 and 2024 — was broadly flat in developing economies: '
     f'{g(S["res_panel"],2021):.1f} percent of adults in 2021 and {g(S["res_panel"],2024):.1f} '
     f'percent in 2024 on the panel; the official aggregate improves modestly, from '
     f'{g(S["res_official"],2021):.1f} to {g(S["res_official"],2024):.1f} percent (Figure 6). '
     'Either way, roughly 45 percent of adults in developing economies could not comfortably '
     'raise emergency funds in 2024. This is the paper’s title in one number: the decade '
     'delivered access at scale, but financial security — the outcome access is meant to enable — '
     'moved on a much slower clock.', keep_next=True)
figure('fig6_resilience.png',
       '*Figure 6.* Could raise emergency funds with some or no difficulty, developing economies, '
       '2021 vs. 2024. Source: author’s calculations from the Global Findex Database 2025.',
       width=3.2)

heading('4.5', 'Who is still left out', level=2)
para('The global gender gap in account ownership halved over the period, from '
     f'{g(S["gender_gap"],2011):.1f} points in 2011 to {g(S["gender_gap"],2024):.1f} points in '
     '2024 (Figure 7), with the sharpest narrowing after 2017 as mobile money and mass '
     'account-opening programmes reached women outside the banking system.', keep_next=True)
figure('fig7_gender.png', width=5.2,
       caption_text='*Figure 7.* Account ownership by gender (left) and the gender gap in percentage points '
       '(right), world. Source: author’s calculations from the Global Findex Database 2025.')
para('The income gap — richest 60 versus poorest 40 percent of households — is consistently '
     f'wider: {g(S["income_gap"],2011):.1f} points in 2011, {g(S["income_gap"],2024):.1f} in 2024 '
     '(Figure 8). Its path is not monotonic: the unusually narrow 2021 gap '
     f'({g(S["income_gap"],2021):.1f} points) coincides with the COVID-era wave, when emergency '
     'government transfers pushed millions of poorer adults into first-time account ownership; '
     'part of that convergence receded by 2024. Economic status, not gender, remains the '
     'strongest observable predictor of exclusion.', keep_next=True)
figure('fig8_income.png', width=4.7, caption_text=
       '*Figure 8.* Account ownership by household income group, world. Annotations show the gap '
       'in percentage points. Source: author’s calculations from the Global Findex Database '
       '2025.')

# ================================================================ 5 PITFALLS
h5 = heading(5, 'How aggregation pitfalls flip trends')
h5.paragraph_format.page_break_before = True
para('Each pitfall of Section 3.4 is now quantified in isolation — holding everything else '
     'correct and toggling one choice — and then in a plausible combination. Table 3 collects '
     'the results.')
make_table(
    ['Pitfall (indicator used for the demonstration)', 'Naive estimate', 'Corrected (panel)',
     'Official aggregate'],
    [
        ['P1 Disaggregation-row leakage\n(emergency funds, developing, 2021→2024)',
         f'{g(S["P1_naive"],2021):.1f} → {g(S["P1_naive"],2024):.1f}',
         f'{g(S["P1_correct"],2021):.1f} → {g(S["P1_correct"],2024):.1f}',
         f'{g(S["res_official"],2021):.1f} → {g(S["res_official"],2024):.1f}'],
        ['P2 Unbalanced composition\n(borrowed formally, developing, 2021→2024)',
         f'{g(S["P2_naive"],2021):.1f} → {g(S["P2_naive"],2024):.1f}',
         f'{g(S["P2_correct"],2021):.1f} → {g(S["P2_correct"],2024):.1f}',
         f'{g(S["bor_dev_official"],2021):.1f} → {g(S["bor_dev_official"],2024):.1f}'],
        ['P3 Indicator-variant confusion\n(saving, developing, 2021→2024)',
         f'{g(S["P3_naive"],2021):.1f} → {g(S["P3_naive"],2024):.1f}',
         f'{g(S["P3_correct"],2021):.1f} → {g(S["P3_correct"],2024):.1f}',
         f'{g(S["sav_dev_official"],2021):.1f} → {g(S["sav_dev_official"],2024):.1f}'],
        ['P4 Coverage-driven missingness\n(mobile money, world, 2024)',
         f'{S["P4"]["naive_2024"]:.1f}', f'{S["P4"]["filled_2024"]:.1f}',
         f'{S["P4"]["official_2024"]:.1f}'],
        ['Compound: P1+P2+P4 with global scope\n(emergency funds, 2021→2024)',
         f'{g(S["orig_resilience"],2021):.1f} → {g(S["orig_resilience"],2024):.1f}',
         f'{g(S["P1_correct"],2021):.1f} → {g(S["P1_correct"],2024):.1f}',
         f'{g(S["res_official"],2021):.1f} → {g(S["res_official"],2024):.1f}'],
        ['Compound: P2+P3 with global scope\n(borrowing, 2021→2024)',
         f'{g(S["orig_borrowing"],2021):.1f} → {g(S["orig_borrowing"],2024):.1f}',
         f'{g(S["P2_correct"],2021):.1f} → {g(S["P2_correct"],2024):.1f}',
         f'{g(S["bor_dev_official"],2021):.1f} → {g(S["bor_dev_official"],2024):.1f}'],
    ],
    col_widths=[2.9, 1.35, 1.3, 1.25], cell_size=9.5,
    caption_above='*Table 3.* Aggregation pitfalls quantified, percent of adults. Isolated '
                  'pitfalls (P1–P4) toggle one choice; compound rows reproduce plausible naive '
                  'pipelines. Naive borrowing uses the narrow at-a-financial-institution variant.')
para('Read individually, the pitfalls look survivable. P1 biases the resilience level by about '
     '1.5 points, because subgroup rows carry the full country weight and minority subgroups with '
     'lower financial inclusion are overrepresented relative to their population shares. P2 '
     'shifts the borrowing series by roughly a point. P3 understates 2024 saving by 4.3 points — '
     'exactly the mobile-money component of the headline definition. P4 is the largest in '
     f'isolation: averaging only the economies that survey mobile money puts the global figure at '
     f'{S["P4"]["naive_2024"]:.1f} percent, nearly double the official '
     f'{S["P4"]["official_2024"]:.1f} percent, because the surveyed subset is dominated by '
     'high-adoption African economies.')
para('The compound rows show why these pitfalls matter beyond level accuracy. Combining the '
     'group-row leak with global scope and wave-varying coverage turns flat resilience '
     f'({g(S["P1_correct"],2021):.1f} → {g(S["P1_correct"],2024):.1f}) into an apparent '
     f'four-and-a-half-point decline ({g(S["orig_resilience"],2021):.1f} → '
     f'{g(S["orig_resilience"],2024):.1f}) — a spurious finding that invites a ready-made '
     'macroeconomic narrative about post-pandemic erosion of household security. Combining '
     'composition drift with the narrow borrowing variant turns gradual growth into an apparent '
     f'six-point fall ({g(S["orig_borrowing"],2021):.1f} → '
     f'{g(S["orig_borrowing"],2024):.1f}): the 2021 full sample contains far more high-income, '
     'high-borrowing economies than 2024, so the composite series tracks sample composition, not '
     'behaviour. A similar mechanism produces an apparent collapse of high-income digital-payment '
     f'usage from {g(S["orig_hi_dp"],2021):.1f} to {g(S["orig_hi_dp"],2024):.1f} percent when the '
     'five reporting economies of 2024 are plotted as if they were the group.')
para('None of these errors survives contact with the file’s own aggregate rows: each naive '
     'series deviates from its official counterpart by amounts far exceeding the panel’s '
     f'{S["account_maxdev"]}-point benchmark tolerance, and the resilience and borrowing '
     'reversals contradict the official direction outright. The practical recommendation is '
     'correspondingly simple, and it is a workflow rather than a technique: fix the group slice '
     'first; fix the country composition; check which indicator variant the concept requires; '
     'count reporting economies before averaging them; and benchmark every series against the '
     'embedded official aggregates before interpreting it.')

# ================================================================ 6 DISCUSSION
heading(6, 'Discussion and limitations')
para('The findings are descriptive. Findex is a repeated cross-section of country-level survey '
     'aggregates; nothing here identifies the causal effect of access on saving, resilience or '
     'anything else. The balanced panel trades a small coverage loss (3.5 percent of the surveyed '
     'adult population) for compositional consistency, and its deviation from the official, '
     'imputation-based aggregates is bounded at about 1.5 points on the headline indicator, with '
     'a known direction. The resilience comparison rests on two waves and a survey question whose '
     'framing invites reporting effects; it should be read as “no evidence of improvement at '
     'scale” rather than as a precise level. Mobile-money estimates treat non-surveyed '
     'economies as zero-adoption; this reproduces the official aggregates but by construction '
     'understates whatever marginal adoption exists in non-surveyed markets. Finally, survey '
     'mode changed for some economies between 2017 and 2021 (face-to-face to telephone), a '
     'comparability caveat the World Bank documents and which applies equally to the official '
     'series.')

# ================================================================ 7 CONCLUSION
heading(7, 'Conclusion')
para('Between 2011 and 2024 the world put an account in the hands of nearly two billion adults — '
     'the fastest documented expansion of access to a formal financial service. The 2024 wave '
     'adds the first genuine usage breakthrough beyond payments: a fourteen-to-sixteen point '
     'surge in formal saving in developing economies, carried largely by mobile money. But the '
     'dimension that access was meant to secure — the ability to absorb a financial shock — '
     'stayed near 55 percent, and the income gap in ownership remains wider than the gender gap. '
     'The agenda implied by the data has shifted from opening accounts to deepening their use.')
para('Methodologically, the paper’s message is that the Findex 2025 file is both a hazard '
     'and its own remedy. Its stacked disaggregations, embedded aggregates and uneven coverage '
     'make sign-flipping aggregation errors easy to produce from individually small mistakes — '
     'and its embedded official aggregates make every one of them detectable in a single '
     'comparison. Validation against the publisher’s own numbers should be the first step '
     'of any analysis built on this database, not an afterthought.')

# ================================================================ back matter
heading(None, 'Data and code availability')
para('The Global Findex Database 2025 is published by the World Bank at '
     'worldbank.org/globalfindex. All code producing every figure, table and statistic in this '
     'paper, together with the analysis notebook, is available at [repository URL].')
heading(None, 'References')
refs = [
    'Demirgüç-Kunt, A., and L. Klapper (2012). “Measuring Financial Inclusion: '
    'The Global Findex Database.” World Bank Policy Research Working Paper 6025. '
    'Washington, DC: World Bank. https://doi.org/10.1596/1813-9450-6025.',
    'Demirgüç-Kunt, A., L. Klapper, D. Singer, and S. Ansar (2022). *The Global Findex '
    'Database 2021: Financial Inclusion, Digital Payments, and Resilience in the Age of '
    'COVID-19.* Washington, DC: World Bank. https://doi.org/10.1596/978-1-4648-1897-4.',
    'Herndon, T., M. Ash, and R. Pollin (2014). “Does High Public Debt Consistently Stifle '
    'Economic Growth? A Critique of Reinhart and Rogoff.” *Cambridge Journal of Economics* '
    '38(2): 257–279. https://doi.org/10.1093/cje/bet075.',
    'Klapper, L., D. Singer, L. Starita, and A. Norris (2025). *The Global Findex Database 2025: '
    'Connectivity and Financial Inclusion in the Digital Economy.* Washington, DC: World Bank. '
    'https://doi.org/10.1596/978-1-4648-2204-9.',
    'Suri, T., and W. Jack (2016). “The Long-Run Poverty and Gender Impacts of Mobile '
    'Money.” *Science* 354(6317): 1288–1292. https://doi.org/10.1126/science.aah5309.',
    'World Bank (2025). *Global Findex Database 2025: Survey Methodology.* Washington, DC: '
    'World Bank. https://www.worldbank.org/en/publication/globalfindex/methodology.',
]
for r in refs:
    p = para(r, size=11, align='left', space_after=6)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)

out = os.path.join(HERE, 'findex_2025_working_paper.docx')
doc.save(out)

# python-docx's default template omits the required w:percent attribute on w:zoom
import zipfile
import shutil
tmp = out + '.tmp'
with zipfile.ZipFile(out) as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
    for item in zin.namelist():
        data = zin.read(item)
        if item == 'word/settings.xml':
            data = data.replace(b'<w:zoom/>', b'<w:zoom w:percent="100"/>')
        zout.writestr(item, data)
shutil.move(tmp, out)
print('saved:', out)
